from fastapi import FastAPI, HTTPException, BackgroundTasks
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from pydantic import BaseModel, HttpUrl
import os
import asyncio
from typing import List, Optional
import tempfile
import uuid
from datetime import datetime
from dotenv import load_dotenv
import re
from urllib.parse import urlparse

# Load environment variables
load_dotenv()

from services.transcript_service import TranscriptService
from services.news_generator import NewsGenerator
from models.response_models import NewsArticle, TranscriptResponse

app = FastAPI(title="Podcast2News API", version="1.0.0")

from fastapi import FastAPI
from fastapi.middleware.cors import CORSMiddleware

app = FastAPI()

# Get allowed origins from environment variable
allowed_origins = os.getenv("ALLOWED_ORIGINS", "http://localhost:3000,http://127.0.0.1:3000").split(",")

# CORS configuration with environment-based origins
app.add_middleware(
    CORSMiddleware,
    allow_origins=allowed_origins,
    allow_credentials=True,
    allow_methods=["GET", "POST"],  # Only allow necessary methods
    allow_headers=["*"],
)

# Initialize services
transcript_service = TranscriptService()
news_generator = NewsGenerator()

# Store generated articles temporarily
articles_store = {}

def validate_youtube_url(url: str) -> bool:
    """Validate if the URL is a valid YouTube URL"""
    youtube_patterns = [
        r'^https?://(www\.)?youtube\.com/watch\?v=[\w-]+',
        r'^https?://(www\.)?youtu\.be/[\w-]+',
        r'^https?://(www\.)?youtube\.com/embed/[\w-]+',
    ]
    return any(re.match(pattern, url) for pattern in youtube_patterns)

class YouTubeURLRequest(BaseModel):
    url: str
    
    def __post_init__(self):
        if not validate_youtube_url(self.url):
            raise ValueError("Invalid YouTube URL format")

class GenerateNewsRequest(BaseModel):
    url: str
    
    def __post_init__(self):
        if not validate_youtube_url(self.url):
            raise ValueError("Invalid YouTube URL format")

@app.get("/")
async def root():
    return {"message": "Podcast2News API is running"}

@app.post("/extract-transcript")
async def extract_transcript(request: YouTubeURLRequest):
    """Extract transcript from YouTube URL"""
    try:
        # Validate YouTube URL
        if not validate_youtube_url(request.url):
            raise HTTPException(status_code=400, detail="Invalid YouTube URL format")
        
        transcript = await transcript_service.get_transcript(request.url)
        return TranscriptResponse(
            success=True,
            transcript=transcript,
            url=request.url
        )
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))

@app.post("/generate-news")
async def generate_news(request: GenerateNewsRequest):
    """Generate news articles from podcast transcript"""
    try:
        # Validate YouTube URL
        if not validate_youtube_url(request.url):
            raise HTTPException(status_code=400, detail="Invalid YouTube URL format")
        
        # Extract transcript first
        transcript = await transcript_service.get_transcript(request.url)
        
        # Generate news articles
        articles = await news_generator.generate_articles(transcript)
        
        # Store articles with unique ID
        session_id = str(uuid.uuid4())
        articles_store[session_id] = articles
        
        return {
            "success": True,
            "session_id": session_id,
            "articles": articles,
            "url": request.url
        }
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))

@app.get("/download-article/{session_id}/{article_index}/{format}")
async def download_article(session_id: str, article_index: int, format: str = "txt"):
    """Download a specific article as a .txt or .docx file"""
    if session_id not in articles_store:
        raise HTTPException(status_code=404, detail="Session not found")
    
    articles = articles_store[session_id]
    if article_index >= len(articles):
        raise HTTPException(status_code=404, detail="Article not found")
    
    if format not in ["txt", "docx"]:
        raise HTTPException(status_code=400, detail="Format must be 'txt' or 'docx'")
    
    article = articles[article_index]
    # Sanitize filename more thoroughly
    safe_title = re.sub(r'[^\w\s-]', '', article.title).strip()
    safe_title = re.sub(r'[-\s]+', '-', safe_title)
    
    if format == "txt":
        # Create temporary txt file
        with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False) as tmp_file:
            tmp_file.write(f"Title: {article.title}\n\n")
            tmp_file.write(f"Content: {article.content}\n\n")
            tmp_file.write(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            tmp_file_path = tmp_file.name
        
        def cleanup_file():
            try:
                os.unlink(tmp_file_path)
            except:
                pass
        
        return FileResponse(
            tmp_file_path,
            media_type='text/plain',
            filename=f"{safe_title}.txt",
            background=BackgroundTasks().add_task(cleanup_file)
        )
    
    else:  # docx format
        from docx import Document
        
        # Create temporary docx file
        doc = Document()
        doc.add_heading(article.title, 0)
        doc.add_paragraph(article.content)
        
        if article.key_quote:
            doc.add_paragraph()
            quote_p = doc.add_paragraph()
            quote_p.add_run(f'"{article.key_quote}"').italic = True
        
        doc.add_paragraph()
        doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_file:
            doc.save(tmp_file.name)
            tmp_file_path = tmp_file.name
        
        def cleanup_file():
            try:
                os.unlink(tmp_file_path)
            except:
                pass
        
        return FileResponse(
            tmp_file_path,
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            filename=f"{safe_title}.docx",
            background=BackgroundTasks().add_task(cleanup_file)
        )

@app.on_event("startup")
async def startup_event():
    print("Podcast2News API started successfully!")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)